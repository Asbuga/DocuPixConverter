from typing import Union

from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfReader


def _read_docx(file_docs: Union[str, bytes]) -> str:
    """Reads DOCX file and return only text content.

    Args:
        filename (Union[str, bytes]): filename or file path

    Returns:
        str: file content
    """
    document = Document(file_docs)
    return "\n".join([line.text for line in document.paragraphs])


def write_docx(
        file_docs: Union[str, bytes], 
        file_pdf: Union[str, bytes]
    ) -> None:
    """Writes textual content to a DOCX file.

    Args:
        file_docs (Union[str, bytes]): filename or file path
        file_pdf (Union[str, bytes]): filename or file path
    """
    content=_read_pdf(file_pdf)
    document = Document()
    for line in content.split("\n"):
        document.add_paragraph(line)
    document.save(file_docs)


def _read_pdf(file_pdf: Union[str, bytes]) -> str:
    """Reads DOCX file and return only text content.

    Args:
        filename (Union[str, bytes]): filename or file path

    Returns:
        str: file content
    """
    reader = PdfReader(stream=file_pdf)
    return "\n".join(
        [reader.pages[i].extract_text() for i in range(0, len(reader.pages))]
    )


def write_pdf(
        file_pdf: Union[str, bytes], 
        file_word: Union[str, bytes]
    ) -> None:
    """Writes textual content to a PDF file.

    Args:
        file_pdf (Union[str, bytes]): filename or file path
        file_word (Union[str, bytes]): filename or file path
    """
    content = _read_docx(file_word)
    pdf = FPDF()
    pdf.set_font("Arial", size=12)
    pdf.add_page()
    pdf.write(5, content)
    pdf.output(file_pdf)
